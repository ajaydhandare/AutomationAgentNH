"""Renders doc/architect.md to doc/architect.docx, for when someone needs a Word copy.

    pip install python-docx
    python doc/make_architect_docx.py [source.md] [target.docx]

architect.md is the source of truth and is the file to edit; this script only formats it, so the
two can never drift. The .docx is a generated artifact and is git-ignored.

Supports the subset of Markdown architect.md actually uses: ATX headings, pipe tables, fenced code
blocks, `-` and `1.` lists, blockquotes, `---` rules, and inline **bold** / *italic* / `code`.
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ACCENT = RGBColor(0x1F, 0x3A, 0x5F)
MUTED = RGBColor(0x5A, 0x5A, 0x5A)
CODE_FILL = "F2F4F7"

INLINE = re.compile(r"(\*\*.+?\*\*|(?<!\*)\*[^*]+?\*(?!\*)|`[^`]+`)", re.DOTALL)
TABLE_DIVIDER = re.compile(r"^\|[\s:|-]+\|$")


def add_runs(paragraph, text):
    """Writes text into a paragraph, honouring **bold**, *italic* and `code`."""
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        elif part.startswith("*") and part.endswith("*"):
            paragraph.add_run(part[1:-1]).italic = True
        else:
            paragraph.add_run(part)


def shade(paragraph, fill):
    pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    pr.append(shd)


def build_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for level, size in ((1, 16), (2, 13), (3, 11.5)):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = ACCENT
        style.font.bold = True
        style.paragraph_format.space_before = Pt(16 if level == 1 else 12)
        style.paragraph_format.space_after = Pt(6)

    section = doc.sections[0]
    section.left_margin = section.right_margin = Inches(0.9)
    section.top_margin = section.bottom_margin = Inches(0.8)

    footer = section.footer.paragraphs[0]
    footer.text = "NewHorizon Automation Agent — Architecture"
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = MUTED


def split_row(line):
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def emit_table(doc, rows):
    """First row is the header. A header of empty strings means a plain two-column fact box."""
    headers, body = rows[0], rows[1:]
    headless = not any(headers)

    table = doc.add_table(rows=0, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    if not headless:
        cells = table.add_row().cells
        for i, header in enumerate(headers):
            paragraph = cells[i].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(2)
            run = paragraph.add_run(header)
            run.bold = True
            run.font.size = Pt(9.5)

    for row in body:
        cells = table.add_row().cells
        for i, value in enumerate(row[: len(headers)]):
            paragraph = cells[i].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(2)
            add_runs(paragraph, value)
            for run in paragraph.runs:
                if run.font.size is None:
                    run.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def emit_code(doc, lines):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    shade(paragraph, CODE_FILL)


def convert(markdown, doc):
    lines = markdown.splitlines()
    index = 0
    title_done = False

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        # fenced code
        if stripped.startswith("```"):
            index += 1
            block = []
            while index < len(lines) and not lines[index].strip().startswith("```"):
                block.append(lines[index])
                index += 1
            emit_code(doc, block)
            index += 1
            continue

        # table: a pipe row followed by the |---| divider
        if (
            stripped.startswith("|")
            and index + 1 < len(lines)
            and TABLE_DIVIDER.match(lines[index + 1].strip())
        ):
            rows = [split_row(stripped)]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                rows.append(split_row(lines[index].strip()))
                index += 1
            emit_table(doc, rows)
            continue

        if not stripped or stripped == "---":
            index += 1
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped[level:].strip()

            if level == 1 and not title_done:
                # The document title, styled rather than treated as a section heading.
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(text)
                run.bold = True
                run.font.size = Pt(24)
                run.font.color.rgb = ACCENT
                paragraph.paragraph_format.space_after = Pt(14)
                title_done = True
            else:
                heading = doc.add_heading(level=min(level - 1, 3) if title_done else level)
                add_runs(heading, text)
                for run in heading.runs:
                    run.font.color.rgb = ACCENT
                    run.bold = True
            index += 1
            continue

        if stripped.startswith("> "):
            quote = doc.add_paragraph()
            quote.paragraph_format.left_indent = Inches(0.25)
            add_runs(quote, stripped[2:])
            for run in quote.runs:
                run.italic = True
                run.font.color.rgb = MUTED
            index += 1
            continue

        if stripped.startswith("- ") or re.match(r"^\d+\.\s", stripped):
            style = "List Bullet" if stripped.startswith("- ") else "List Number"
            text = stripped[2:] if stripped.startswith("- ") else re.sub(r"^\d+\.\s", "", stripped)
            # a wrapped list item continues on indented following lines
            index += 1
            while index < len(lines) and lines[index].startswith("  ") and lines[index].strip():
                text += " " + lines[index].strip()
                index += 1
            item = doc.add_paragraph(style=style)
            item.paragraph_format.space_after = Pt(3)
            add_runs(item, text)
            continue

        # a paragraph runs until a blank line
        block = [stripped]
        index += 1
        while index < len(lines) and lines[index].strip() and not lines[index].strip().startswith(
            ("#", "|", "```", "- ", "> ")
        ):
            block.append(lines[index].strip())
            index += 1
        add_runs(doc.add_paragraph(), " ".join(block))


def main():
    here = Path(__file__).resolve().parent
    source = Path(sys.argv[1]) if len(sys.argv) > 1 else here / "architect.md"
    target = Path(sys.argv[2]) if len(sys.argv) > 2 else here / "architect.docx"

    doc = Document()
    build_styles(doc)
    convert(source.read_text(encoding="utf-8"), doc)
    doc.save(target)

    print(f"written: {target}")


if __name__ == "__main__":
    main()
