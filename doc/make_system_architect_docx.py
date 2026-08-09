"""Renders doc/System Architect.md — figures and all — to doc/System Architect.docx.

    pip install python-docx pillow
    python doc/make_system_architect_docx.py

The markdown file is the source of truth and is the file to edit; the figures are drawn by
system_architect_diagrams.py. This script only formats them, so the three can never drift.

Beyond the markdown subset that make_architect_docx.py handles, this one adds what a stakeholder
document needs: a cover page, `![caption](path)` figures with numbered captions, `<!-- pagebreak -->`,
and page numbers in the footer.
"""

import importlib.util
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Inches, Pt, RGBColor
from PIL import Image

HERE = Path(__file__).resolve().parent

ACCENT = RGBColor(0x1F, 0x3A, 0x5F)
MUTED = RGBColor(0x5A, 0x5A, 0x5A)

IMAGE = re.compile(r"^!\[(?P<caption>.*)\]\((?P<path>[^)]+)\)$")
PAGEBREAK = "<!-- pagebreak -->"
FIGURE_WIDTH = Inches(6.5)


def _shared():
    """Reuses the inline/table/code primitives from the sibling renderer rather than copying them."""
    spec = importlib.util.spec_from_file_location(
        "make_architect_docx", HERE / "make_architect_docx.py"
    )
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


SHARED = _shared()
add_runs = SHARED.add_runs
emit_code = SHARED.emit_code
emit_table = SHARED.emit_table
split_row = SHARED.split_row
TABLE_DIVIDER = SHARED.TABLE_DIVIDER


def build_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for level, size in ((1, 16), (2, 13), (3, 11.5)):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = ACCENT
        style.font.bold = True
        style.paragraph_format.space_before = Pt(16 if level == 1 else 12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True

    section = doc.sections[0]
    section.left_margin = section.right_margin = Inches(0.9)
    section.top_margin = section.bottom_margin = Inches(0.8)


def _field(paragraph, instruction):
    """Inserts a Word field code — used for the page number, which Word must compute itself."""
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run = paragraph.add_run()
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED
    for element in (begin, instr, end):
        run._r.append(element)


def build_footer(section, title):
    footer = section.footer.paragraphs[0]
    footer.text = ""
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = footer.add_run(f"{title}    |    ")
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED
    _field(footer, " PAGE ")
    run = footer.add_run(" of ")
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED
    _field(footer, " NUMPAGES ")


def emit_cover(doc, title, subtitle):
    for _ in range(4):
        doc.add_paragraph()

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(title)
    run.bold = True
    run.font.size = Pt(30)
    run.font.color.rgb = ACCENT

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(20)
    run = paragraph.add_run(subtitle)
    run.font.size = Pt(15)
    run.font.color.rgb = MUTED

    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = rule.add_run("—" * 18)
    run.font.color.rgb = ACCENT


def emit_figure(doc, path, caption, number):
    """Places a figure at page width, with a numbered caption beneath it."""
    if not path.exists():
        raise FileNotFoundError(
            f"{path} is missing — run 'python doc/system_architect_diagrams.py' first."
        )

    with Image.open(path) as image:
        width, height = image.size

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.add_run().add_picture(
        str(path), width=FIGURE_WIDTH, height=Emu(int(FIGURE_WIDTH * height / width))
    )

    # The drawn caption inside the image says "Figure N"; this one is the searchable Word caption.
    text = caption.split("—", 1)[-1].strip() if "—" in caption else caption
    label = doc.add_paragraph()
    label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label.paragraph_format.space_after = Pt(12)
    run = label.add_run(f"Figure {number} — ")
    run.bold = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = ACCENT
    run = label.add_run(text)
    run.italic = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = MUTED


def convert(markdown, doc, base_dir):
    lines = markdown.splitlines()
    index = 0
    title_done = False
    figure_number = 0

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if stripped == PAGEBREAK:
            doc.add_page_break()
            index += 1
            continue

        match = IMAGE.match(stripped)
        if match:
            figure_number += 1
            emit_figure(
                doc,
                (base_dir / match.group("path")).resolve(),
                match.group("caption"),
                figure_number,
            )
            index += 1
            continue

        if stripped.startswith("```"):
            index += 1
            block = []
            while index < len(lines) and not lines[index].strip().startswith("```"):
                block.append(lines[index])
                index += 1
            emit_code(doc, block)
            index += 1
            continue

        if (
            stripped.startswith("|")
            and index + 1 < len(lines)
            and TABLE_DIVIDER.match(lines[index + 1].strip())
        ):
            rows = [split_row(stripped)]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                rows.append(split_row(lines[index].strip()))
                index += 1
            emit_table(doc, rows)
            continue

        if not stripped or stripped == "---":
            index += 1
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped[level:].strip()

            if level == 1 and not title_done:
                title_done = True
                cover_title = text
                # The document's H2 immediately after the H1 is the cover subtitle.
                cover_subtitle = ""
                if index + 1 < len(lines) and lines[index + 1].strip().startswith("## "):
                    cover_subtitle = lines[index + 1].strip()[3:].strip()
                    index += 1
                emit_cover(doc, cover_title, cover_subtitle)
                index += 1
                continue

            heading = doc.add_heading(level=min(level - 1, 3))
            add_runs(heading, text)
            for run in heading.runs:
                run.font.color.rgb = ACCENT
                run.bold = True
            index += 1
            continue

        if stripped.startswith("> "):
            quote = doc.add_paragraph()
            quote.paragraph_format.left_indent = Inches(0.25)
            quote.paragraph_format.right_indent = Inches(0.25)
            block = [stripped[2:]]
            index += 1
            while index < len(lines) and lines[index].strip().startswith(">"):
                block.append(lines[index].strip().lstrip(">").strip())
                index += 1
            add_runs(quote, " ".join(part for part in block if part))
            for run in quote.runs:
                run.italic = True
                run.font.color.rgb = MUTED
            SHARED.shade(quote, "F5F7FA")
            continue

        if stripped.startswith("- ") or re.match(r"^\d+\.\s", stripped):
            style = "List Bullet" if stripped.startswith("- ") else "List Number"
            text = stripped[2:] if stripped.startswith("- ") else re.sub(r"^\d+\.\s", "", stripped)
            index += 1
            while index < len(lines) and lines[index].startswith("  ") and lines[index].strip():
                text += " " + lines[index].strip()
                index += 1
            item = doc.add_paragraph(style=style)
            item.paragraph_format.space_after = Pt(3)
            add_runs(item, text)
            continue

        block = [stripped]
        index += 1
        while index < len(lines) and lines[index].strip() and not lines[index].strip().startswith(
            ("#", "|", "```", "- ", "> ", "![", "<!--")
        ):
            block.append(lines[index].strip())
            index += 1
        add_runs(doc.add_paragraph(), " ".join(block))

    return figure_number


def main():
    source = Path(sys.argv[1]) if len(sys.argv) > 1 else HERE / "System Architect.md"
    target = Path(sys.argv[2]) if len(sys.argv) > 2 else HERE / "System Architect.docx"

    doc = Document()
    build_styles(doc)
    build_footer(doc.sections[0], "NewHorizon Automation Agent — System Architecture")
    figures = convert(source.read_text(encoding="utf-8"), doc, source.parent)
    doc.core_properties.title = "NewHorizon Automation Agent — System Architecture"
    doc.core_properties.subject = "System architecture, technology rationale and data model"
    doc.save(target)

    print(f"written: {target}  ({figures} figures)")


if __name__ == "__main__":
    main()
